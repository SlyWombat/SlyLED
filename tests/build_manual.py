#!/usr/bin/env python3
"""
build_manual.py — Generate SlyLED User Manual as a Word document.

Reads docs/USER_MANUAL.md, inserts screenshots from docs/screenshots/,
and outputs docs/SlyLED_User_Manual.docx.

Run screenshot_capture.py first to ensure screenshots are up to date.

Usage:
    python tests/build_manual.py                    # Build manual
    python tests/build_manual.py --screenshots      # Regenerate screenshots first
"""

import sys, os, re, argparse, subprocess
from docx.shared import RGBColor

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'desktop', 'shared'))

PROJ = os.path.join(os.path.dirname(__file__), '..')
DOCS = os.path.join(PROJ, 'docs')
SHOTS = os.path.join(DOCS, 'screenshots')
MANUAL_MD = os.path.join(DOCS, 'USER_MANUAL.md')
OUTPUT = os.path.join(DOCS, 'SlyLED_User_Manual.docx')
LOGO = os.path.join(PROJ, 'images', 'slyled.png')

# ── i18n translations ─────────────────────────────────────────────────────
_LANG = 'en'
_TR = {
  'fr': {
    'User Manual': "Manuel d'utilisation",
    '3D Volumetric Lighting System': "Syst\u00e8me d'\u00e9clairage volum\u00e9trique 3D",
    'Complete guide to designing, programming, and running LED + DMX lighting shows with the SlyLED orchestrator, performers, and DMX bridge.':
        "Guide complet pour concevoir, programmer et ex\u00e9cuter des spectacles lumineux LED + DMX avec l'orchestrateur SlyLED, les performers et le pont DMX.",
    'Table of Contents': 'Table des mati\u00e8res',
    '1. Getting Started with 3D Stage Design': '1. D\u00e9marrer avec la conception de sc\u00e8ne 3D',
    'Overview': 'Aper\u00e7u',
    'SlyLED is a three-tier lighting system: the Orchestrator (Windows/Mac desktop app) designs and controls shows, Performers (ESP32/D1 Mini boards with LED strings) execute lighting effects, and DMX Bridges (Giga R1 boards) drive professional DMX fixtures over Art-Net.':
        "SlyLED est un syst\u00e8me d'\u00e9clairage \u00e0 trois niveaux\u00a0: l'Orchestrateur (application Windows/Mac) con\u00e7oit et contr\u00f4le les spectacles, les Performers (cartes ESP32/D1 Mini avec rubans LED) ex\u00e9cutent les effets lumineux, et les Ponts DMX (cartes Giga R1) pilotent les projecteurs DMX professionnels via Art-Net.",
    'Switching Between 2D and 3D': 'Basculer entre 2D et 3D',
    'Navigating the 3D Viewport': 'Naviguer dans la vue 3D',
    'Coordinate System': 'Syst\u00e8me de coordonn\u00e9es',
    '2. Fixture Setup': '2. Configuration des projecteurs',
    'What Are Fixtures?': 'Que sont les projecteurs\u00a0?',
    'Adding an LED Performer': "Ajouter un performer LED",
    'Adding a DMX Fixture': 'Ajouter un projecteur DMX',
    'Editing DMX Fixtures': 'Modifier les projecteurs DMX',
    'Fixture Types': 'Types de projecteurs',
    '3. Creating Spatial Effects': '3. Cr\u00e9er des effets spatiaux',
    'Spatial Fields': 'Champs spatiaux',
    'Moving Heads + Spatial Effects': 'Lyres + effets spatiaux',
    'DMX Action Types': "Types d'actions DMX",
    '4. Building a Timeline': '4. Construire une timeline',
    'Creating a Timeline': 'Cr\u00e9er une timeline',
    '5. Baking & Playback': '5. Compilation et lecture',
    'What Is Baking?': "Qu'est-ce que la compilation\u00a0?",
    'Playback': 'Lecture',
    '6. Show Preview Emulator': "6. \u00c9mulateur d'aper\u00e7u du spectacle",
    '7. Preset Shows': '7. Spectacles pr\u00e9d\u00e9finis',
    '8. DMX Fixture Profiles': '8. Profils de projecteurs DMX',
    'Viewing a Profile': 'Voir un profil',
    'Creating Custom Profiles': 'Cr\u00e9er des profils personnalis\u00e9s',
    'Importing from Open Fixture Library': "Importer depuis l'Open Fixture Library",
    'Import/Export Bundles': 'Import/export de lots',
    '9. Settings & Configuration': '9. Param\u00e8tres et configuration',
    'Firmware': 'Firmware',
    '10. System Limits': '10. Limites du syst\u00e8me',
    '11. Troubleshooting': '11. D\u00e9pannage',
    '12. API Quick Reference': "12. R\u00e9f\u00e9rence rapide de l'API",
    'Resource': 'Ressource', 'Limit': 'Limite', 'Notes': 'Notes',
    'Type': 'Type', 'Description': 'Description', 'Use Case': 'Cas d\'utilisation',
    'Field': 'Champ', 'Action': 'Action', 'Control': 'Contr\u00f4le',
    'Method': 'M\u00e9thode', 'Endpoint': 'Point d\'acc\u00e8s',
    'Preset': 'Pr\u00e9r\u00e9glage',
    'Action Type': "Type d'action",
    'DMX Scene': 'Sc\u00e8ne DMX',
    'Pan/Tilt Move': 'Mouvement Pan/Tilt',
    'Gobo Select': 'S\u00e9lection de gobo',
    'Color Wheel': 'Roue de couleurs',
    # New calibration chapters (#329-#336)
    '10. Camera Setup & Calibration': '10. Configuration des cam\u00e9ras et calibration',
    'Adding Camera Nodes': 'Ajouter des n\u0153uds cam\u00e9ra',
    'ArUco Marker Calibration': "Calibration par marqueurs ArUco",
    'Stage Map — Camera Pose': "Carte de sc\u00e8ne \u2014 Pose de la cam\u00e9ra",
    'Mover Calibration Wizard': "Assistant de calibration des lyres",
    '11. 3D Environment Scanning': '11. Scan 3D de l\u2019environnement',
    'Point Cloud Capture': 'Capture du nuage de points',
    'Surface Detection': 'D\u00e9tection des surfaces',
    'CV Engine': 'Moteur de vision par ordinateur',
    '12. Stereo 3D & Light Mapping': '12. St\u00e9r\u00e9o 3D et cartographie lumineuse',
    'Stereo Triangulation': 'Triangulation st\u00e9r\u00e9o',
    'Per-Fixture Light Maps': 'Cartographie lumineuse par projecteur',
    'Beam-as-Structured-Light': 'Faisceau comme lumi\u00e8re structur\u00e9e',
    '13. Project Files': '13. Fichiers de projet',
    'Export': 'Exportation',
    'Import': 'Importation',
    'Spatial Data in Projects': 'Donn\u00e9es spatiales dans les projets',
    '14. System Limits': '14. Limites du syst\u00e8me',
    '15. Troubleshooting': '15. D\u00e9pannage',
    '16. API Quick Reference': "16. R\u00e9f\u00e9rence rapide de l'API",
  }
}
def T(s):
    """Translate string s using current language."""
    if _LANG == 'en':
        return s
    return _TR.get(_LANG, {}).get(s, s)

# Kinetic Prism design palette
CLR_BG = RGBColor(0x0A, 0x0F, 0x13) if False else None  # can't set page bg in python-docx
CLR_HEADING = RGBColor(0x38, 0xBD, 0xF8)   # LuminaBlue
CLR_ACCENT = RGBColor(0x22, 0xD3, 0xEE)    # Cyan accent
CLR_TEXT = RGBColor(0x33, 0x33, 0x33)        # Dark text (print-friendly)
CLR_MUTED = RGBColor(0x64, 0x74, 0x8B)      # Muted gray
CLR_PURPLE = RGBColor(0x7C, 0x3A, 0xED)     # DMX purple
CLR_GREEN = RGBColor(0x22, 0xC5, 0x5E)      # LED green


def build_manual():
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # ── Styles — Kinetic Prism theme ──────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = CLR_TEXT

    # Heading 1: LuminaBlue
    h1 = doc.styles['Heading 1']
    h1.font.color.rgb = CLR_HEADING
    h1.font.name = 'Calibri'
    h1.font.size = Pt(22)
    h1.font.bold = True

    # Heading 2: Purple accent
    h2 = doc.styles['Heading 2']
    h2.font.color.rgb = CLR_PURPLE
    h2.font.name = 'Calibri'
    h2.font.size = Pt(15)

    # Heading 3: Cyan accent
    h3 = doc.styles['Heading 3']
    h3.font.color.rgb = CLR_ACCENT
    h3.font.name = 'Calibri'
    h3.font.size = Pt(13)

    # ── Screenshot mapping: section keyword → screenshot files ──
    SECTION_SHOTS = {
        '1. Getting Started': [
            ('spa-layout-2d.png', 'Layout tab — 2D Canvas view with placed fixtures'),
            ('spa-layout-3d.png', 'Layout tab — 3D Viewport with stage wireframe and fixtures'),
        ],
        '2. Fixture Setup': [
            ('spa-setup.png', 'Setup tab showing LED and DMX fixtures with type badges'),
            ('spa-setup-add-led.png', 'Add Fixture modal — LED performer flow'),
            ('spa-setup-add-dmx.png', 'Add Fixture modal — DMX fixture with profile dropdown'),
            ('spa-setup-edit-dmx.png', 'Edit DMX fixture — universe, address, aim point, test channels'),
        ],
        '3. Creating Spatial': [
            ('spa-actions.png', 'Actions tab — spatial effects and classic action library'),
        ],
        '4. Building a Timeline': [
            ('spa-runtime.png', 'Runtime tab — timeline editor with tracks and clips'),
        ],
        '5. Baking': [],  # no specific screenshot yet
        '6. Show Preview': [
            ('spa-runtime.png', 'Runtime tab with emulator preview'),
        ],
        '7. Preset Shows': [],  # text-only section
        '10. Camera Setup': [
            ('spa-settings-cameras.png', 'Settings > Cameras — camera list with calibration status'),
        ],
        '11. 3D Environment': [
            ('spa-layout-3d.png', '3D Viewport with point cloud and detected surfaces'),
        ],
        '12. Stereo 3D': [],
        '13. Project': [],
        '14. System Limits': [],
        '15. Troubleshooting': [],
        'Settings': [
            ('spa-settings.png', 'Settings tab — app config and DMX output'),
            ('spa-settings-profiles.png', 'Fixture Profile Library — 12 built-in profiles'),
            ('spa-settings-profile-view.png', 'Profile detail view with channel capabilities'),
            ('spa-settings-profile-editor.png', 'Profile editor — channel table and capability ranges'),
            ('spa-settings-ofl-import.png', 'Open Fixture Library import modal'),
        ],
        'Firmware': [
            ('spa-firmware.png', 'Firmware tab — board detection, version query, flash controls'),
        ],
        'Dashboard': [
            ('spa-dashboard.png', 'Dashboard — live performer status overview'),
        ],
    }

    # ── Read version from version.h ──────────────────────────
    lang_suffix = f'_{_LANG}' if _LANG != 'en' else ''
    version = "8.3"
    vh_path = os.path.join(PROJ, 'main', 'version.h')
    if os.path.exists(vh_path):
        with open(vh_path) as vf:
            vtxt = vf.read()
            import re as _re
            ma = _re.search(r'APP_MAJOR\s+(\d+)', vtxt)
            mi = _re.search(r'APP_MINOR\s+(\d+)', vtxt)
            pa = _re.search(r'APP_PATCH\s+(\d+)', vtxt)
            if ma and mi:
                version = f"{ma.group(1)}.{mi.group(1)}"
                if pa:
                    version += f".{pa.group(1)}"

    # ── Title page — Kinetic Prism styled ─────────────────────
    doc.add_paragraph()

    # Logo
    if os.path.exists(LOGO):
        doc.add_picture(LOGO, width=Inches(3.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    title = doc.add_heading('SlyLED ' + T('User Manual'), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = CLR_HEADING
        run.font.size = Pt(36)

    sub = doc.add_paragraph(f'{T("3D Volumetric Lighting System")} — v{version}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].font.color.rgb = CLR_MUTED

    doc.add_paragraph()
    desc = doc.add_paragraph(T(
        'Complete guide to designing, programming, and running LED + DMX lighting shows '
        'with the SlyLED orchestrator, performers, and DMX bridge.'
    ))
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.runs[0].font.color.rgb = CLR_TEXT

    doc.add_paragraph()
    footer_p = doc.add_paragraph('electricrv.ca/slyled')
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.color.rgb = CLR_ACCENT
    footer_p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ── Table of Contents ──────────────────────────────────────
    doc.add_heading(T('Table of Contents'), level=1)
    toc_items = [
        '1. Getting Started with 3D Stage Design',
        '2. Fixture Setup',
        '3. Creating Spatial Effects',
        '4. Building a Timeline',
        '5. Baking & Playback',
        '6. Show Preview Emulator',
        '7. Preset Shows',
        '8. DMX Fixture Profiles',
        '9. Settings & Configuration',
        '10. Camera Setup & Calibration',
        '11. 3D Environment Scanning',
        '12. Stereo 3D & Light Mapping',
        '13. Project Files',
        '14. System Limits',
        '15. Troubleshooting',
        '16. API Quick Reference',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()

    def add_screenshot(filename, caption):
        """Add a screenshot image with caption."""
        path = os.path.join(SHOTS, filename)
        if not os.path.exists(path):
            doc.add_paragraph(f'[Screenshot: {filename} — not found]').italic = True
            return
        doc.add_picture(path, width=Inches(6.0))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = CLR_MUTED
        cap.runs[0].italic = True

    def add_table(headers, rows):
        """Add a formatted table."""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                table.rows[ri + 1].cells[ci].text = str(val)

    # ── Section 1: Getting Started ─────────────────────────────
    doc.add_heading(T('1. Getting Started with 3D Stage Design'), level=1)

    doc.add_heading(T('Overview'), level=2)
    doc.add_paragraph(
        'SlyLED is a three-tier lighting system: the Orchestrator (Windows/Mac desktop app) '
        'designs and controls shows, Performers (ESP32/D1 Mini boards with LED strings) execute '
        'lighting effects, and DMX Bridges (Giga R1 boards) drive professional DMX fixtures over Art-Net.'
    )

    add_screenshot('spa-dashboard.png', 'Dashboard — live performer status overview')

    doc.add_heading(T('Switching Between 2D and 3D'), level=2)
    doc.add_paragraph(
        'The Layout tab offers two views:\n'
        '- 2D Canvas: Flat top-down layout for simple setups\n'
        '- 3D Viewport: Interactive Three.js scene for complex multi-level installations\n\n'
        'Both views share the same position data. Switching is instant and non-destructive.'
    )

    add_screenshot('spa-layout-2d.png', 'Layout tab — 2D Canvas with placed fixtures')
    add_screenshot('spa-layout-3d.png', 'Layout tab — 3D Viewport with stage wireframe, fixtures, and objects')

    doc.add_heading(T('Navigating the 3D Viewport'), level=2)
    add_table(
        ['Action', 'Control'],
        [
            ['Orbit (rotate)', 'Left-click + drag'],
            ['Zoom', 'Scroll wheel'],
            ['Pan (shift view)', 'Right-click + drag'],
            ['Select fixture', 'Left-click on node'],
            ['Move fixture', 'Drag the colored arrows'],
            ['Edit fixture', 'Double-click on node'],
            ['Place from sidebar', 'Drag unplaced fixture into viewport'],
        ]
    )

    doc.add_heading(T('Coordinate System'), level=2)
    doc.add_paragraph(
        '- X-axis (red): Width — left to right\n'
        '- Y-axis (green): Height — ground to ceiling\n'
        '- Z-axis (blue): Depth — front to back\n'
        '- Origin: Bottom-left-front corner of the stage\n'
        '- Units: Internally millimeters; displayed in Settings-chosen unit'
    )
    doc.add_page_break()

    # ── Section 2: Fixture Setup ───────────────────────────────
    doc.add_heading(T('2. Fixture Setup'), level=1)

    doc.add_heading(T('What Are Fixtures?'), level=2)
    doc.add_paragraph(
        'A fixture is the primary entity on the 3D stage. It wraps physical hardware '
        'and adds stage-level attributes like position, rotation, and aim point.\n\n'
        '- LED Performers: Auto-created when hardware is registered via Setup tab\n'
        '- DMX Fixtures: Created manually with universe, address, and profile assignment\n'
        '- Fixtures can override child attributes (e.g., rotate a horizontal string to vertical)\n'
        '- The baking engine uses the fixture\'s position and rotation, not the child\'s raw config'
    )

    add_screenshot('spa-setup.png', 'Setup tab — LED and DMX fixtures with type badges and status')

    doc.add_heading(T('Adding an LED Performer'), level=2)
    doc.add_paragraph(
        'Click Add Fixture, select "SlyLED Performer (LED)", and enter the device IP address. '
        'The system probes the device via UDP PING and HTTP, registers it as a child, '
        'and auto-creates a linked fixture.'
    )
    add_screenshot('spa-setup-add-led.png', 'Add Fixture — LED performer flow with IP address entry')

    doc.add_heading(T('Adding a DMX Fixture'), level=2)
    doc.add_paragraph(
        'Click Add Fixture, select "DMX Fixture", and configure:\n'
        '- Name: Descriptive label\n'
        '- Universe: DMX universe number (1+)\n'
        '- Start Address: DMX start channel (1-512)\n'
        '- Channel Count: Number of channels the fixture uses\n'
        '- Profile: Select from the built-in library or import from Open Fixture Library\n\n'
        'DMX fixtures appear in the setup table with purple "DMX" badges.'
    )
    add_screenshot('spa-setup-add-dmx.png', 'Add Fixture — DMX fixture with profile dropdown')

    doc.add_heading(T('Editing DMX Fixtures'), level=2)
    doc.add_paragraph(
        'Click Edit on a DMX fixture to modify its properties. The edit modal shows:\n'
        '- Universe, Start Address, Channel Count\n'
        '- Profile ID (for channel name/type mapping)\n'
        '- Aim Point (X, Y, Z in mm) — where the beam points for moving heads\n'
        '- Test Channels — interactive sliders to control each DMX channel live\n'
        '- Rotation Override (degrees)'
    )
    add_screenshot('spa-setup-edit-dmx.png', 'Edit DMX fixture — aim point, profile, and test channels')

    doc.add_heading(T('Fixture Types'), level=2)
    add_table(
        ['Type', 'Description', 'Use Case'],
        [
            ['Linear', 'LED strip/string with pixels along a path', 'LED performers'],
            ['Point', 'Single light source with area of effect', 'DMX pars, spots, moving heads'],
            ['Object', '3D mesh as a projection target', 'Walls, screens'],
            ['Group', 'Named collection of fixtures', 'Zones, grouped control'],
        ]
    )
    doc.add_page_break()

    # ── Section 3: Spatial Effects ─────────────────────────────
    doc.add_heading(T('3. Creating Spatial Effects'), level=1)

    doc.add_paragraph(
        'Spatial effects operate in 3D space. A sphere of colored light sweeping across the stage '
        'illuminates different fixtures at different times based on their physical position. '
        'For DMX moving heads, the effect center becomes the aim target — heads track the effect as it moves.'
    )

    add_screenshot('spa-actions.png', 'Actions tab — spatial effects and classic action library')

    doc.add_heading(T('Spatial Fields'), level=2)
    add_table(
        ['Field', 'Description'],
        [
            ['Shape', 'Sphere, Plane, or Box'],
            ['Color', 'RGB color applied to pixels inside the field'],
            ['Size', 'Radius (sphere), thickness (plane), or W/H/D (box)'],
            ['Motion Start', 'Starting position [x, y, z] in mm'],
            ['Motion End', 'Ending position [x, y, z] in mm'],
            ['Duration', 'Travel time from start to end'],
            ['Easing', 'Linear, ease-in, ease-out, ease-in-out'],
            ['Blend', 'Replace, Add, Multiply, or Screen'],
        ]
    )

    doc.add_heading(T('Moving Heads + Spatial Effects'), level=2)
    doc.add_paragraph(
        'When a spatial effect is applied to a DMX moving head fixture:\n'
        '- The effect\'s center position becomes the aim target\n'
        '- Pan/tilt angles are computed from fixture position to aim point\n'
        '- Color is applied from the effect\'s RGB values\n'
        '- Moving effects produce time-sliced segments that track the motion\n'
        '- 3D viewport shows beam cones pointing at the effect center'
    )

    doc.add_heading(T('DMX Action Types'), level=2)
    doc.add_paragraph(
        'In addition to classic LED actions (Solid, Chase, Rainbow, etc.), four DMX-specific '
        'action types are available for direct control of DMX features:'
    )
    add_table(
        ['Action Type', 'Description'],
        [
            ['DMX Scene', 'Set exact values for dimmer, pan, tilt, strobe, gobo, color wheel, prism'],
            ['Pan/Tilt Move', 'Animate pan/tilt from start to end position over time'],
            ['Gobo Select', 'Select a gobo wheel position with optional color'],
            ['Color Wheel', 'Select a color wheel position'],
        ]
    )
    doc.add_paragraph(
        'Classic LED actions (Solid, Breathe, Chase, etc.) are automatically converted to '
        'DMX Scene segments when assigned to DMX fixtures. The dimmer is auto-set to 255 '
        'when color is active, and pan/tilt default to center (0.5).'
    )
    doc.add_page_break()

    # ── Section 4: Building a Timeline ─────────────────────────
    doc.add_heading(T('4. Building a Timeline'), level=1)

    doc.add_paragraph(
        'Timelines are multi-track, overlapping effect sequences with precise timing. '
        'Use the Runtime tab to create and edit timelines.'
    )

    add_screenshot('spa-runtime.png', 'Runtime tab — timeline editor')

    doc.add_heading(T('Creating a Timeline'), level=2)
    doc.add_paragraph(
        '1. Click "+ New Timeline" and enter name and duration\n'
        '2. Select the timeline from the dropdown\n'
        '3. Add tracks (one per fixture or "All Performers")\n'
        '4. Add clips referencing spatial effects or classic actions\n'
        '5. Adjust clip timing by editing start time and duration'
    )
    doc.add_page_break()

    # ── Section 5: Baking & Playback ──────────────────────────
    doc.add_heading(T('5. Baking & Playback'), level=1)

    doc.add_heading(T('What Is Baking?'), level=2)
    doc.add_paragraph(
        'Baking compiles a timeline into minimal action instructions for each performer. '
        'The smart bake engine analyzes each clip\'s spatial geometry directly and computes '
        'per-string sweep patterns, directions, and speeds.\n\n'
        'For DMX fixtures, baking also computes:\n'
        '- Pan/tilt angles from spatial effect motion paths\n'
        '- Color values from the effect\'s intersection with the fixture position\n'
        '- Dimmer values (auto 255 when color is active)\n'
        '- Time-sliced segments for smooth moving head tracking'
    )

    doc.add_heading(T('Playback'), level=2)
    doc.add_paragraph(
        'After baking and syncing:\n'
        '- LED performers receive action steps via UDP and execute locally\n'
        '- DMX fixtures are driven by a 40Hz playback loop sending Art-Net/sACN packets\n'
        '- All channels (RGB, pan, tilt, dimmer, strobe, gobo) are sent per profile\n'
        '- 16-bit channels (pan/tilt) are split into coarse + fine bytes automatically'
    )
    doc.add_page_break()

    # ── Section 6: Show Preview ───────────────────────────────
    doc.add_heading(T('6. Show Preview Emulator'), level=1)
    doc.add_paragraph(
        'The emulator shows a real-time preview of your show on the Runtime tab.\n\n'
        '- LED fixtures: Colored dots along string directions, animated per-pixel\n'
        '- DMX fixtures: Beam triangles from fixture position toward aim point\n'
        '- Beam color and alpha reflect the fixture\'s current DMX state\n'
        '- Time counter shows elapsed position'
    )
    doc.add_page_break()

    # ── Section 7: Preset Shows ───────────────────────────────
    doc.add_heading(T('7. Preset Shows'), level=1)
    doc.add_paragraph(
        '14 themed shows are available from the Runtime tab. Shows are dynamically generated '
        'based on your actual fixtures — every fixture gets coverage with no dark periods. '
        'LED fixtures get pattern effects, DMX pars get color washes, and moving heads get '
        'pan/tilt sweeps that track spatial effects across the stage.\n\n'
        'Each load produces a unique variation of the theme with randomized timing and positions.'
    )

    add_table(
        ['Preset', 'Type', 'Description'],
        [
            ['Rainbow Up', 'Spatial plane', 'Rainbow from floor to ceiling'],
            ['Rainbow Across', 'Spatial sphere', 'Rainbow sweeping left to right'],
            ['Slow Fire', 'Classic action', 'Warm fire effect on all fixtures'],
            ['Disco', 'Classic action', 'Pastel twinkle sparkles'],
            ['Ocean Wave', 'Spatial (2 effects)', 'Blue wave sweep with teal wash'],
            ['Sunset Glow', 'Mixed', 'Warm breathe with golden plane sweep'],
            ['Police Lights', 'Mixed', 'Red strobe with blue box flash sweep'],
            ['Starfield', 'Classic action', 'White sparkles on dark background'],
            ['Aurora Borealis', 'Spatial (2 effects)', 'Green curtain with purple shimmer'],
            ['Spotlight Sweep', 'Spatial (moving heads)', 'Warm orb sweeps stage — heads track it'],
            ['Concert Wash', 'Mixed (moving heads)', 'Magenta flood + amber tracking spot'],
            ['Figure Eight', 'Spatial (moving heads)', 'Crossing orbs — heads trace X paths'],
            ['Thunderstorm', 'Mixed (moving heads)', 'Lightning strikes — heads chase bolts'],
            ['Dance Floor', 'Mixed (moving heads)', 'Fast orbiting spots — rapid tracking'],
        ]
    )
    doc.add_page_break()

    # ── Section 8: DMX Fixture Profiles ───────────────────────
    doc.add_heading(T('8. DMX Fixture Profiles'), level=1)

    doc.add_paragraph(
        'Profiles define the channel layout and capabilities of DMX fixtures. '
        'Each channel maps a DMX offset to a function (pan, tilt, red, green, blue, etc.) '
        'with capability ranges that describe what each DMX value range does.'
    )

    add_screenshot('spa-settings-profiles.png', 'Fixture Profile Library — 12 built-in profiles')

    doc.add_heading(T('Viewing a Profile'), level=2)
    doc.add_paragraph(
        'Click View on any profile to see the full channel table with capabilities. '
        'Moving head profiles show pan/tilt ranges, beam width, and color mode.'
    )
    add_screenshot('spa-settings-profile-view.png', 'Moving Head 16-bit profile with 13 channels and capability ranges')

    doc.add_heading(T('Creating Custom Profiles'), level=2)
    doc.add_paragraph(
        'Click "New Profile" to open the editor. Define channels with type, name, bits (8/16), '
        'and capability ranges. Click the capabilities button on each channel to define '
        'what DMX value ranges do (e.g., gobo positions, strobe speeds, color wheel slots).'
    )
    add_screenshot('spa-settings-profile-editor.png', 'Profile editor — channel table with type dropdowns')
    add_screenshot('workflow-profile-caps.png', 'Capability range editor for a channel')

    doc.add_heading(T('Importing from Open Fixture Library'), level=2)
    doc.add_paragraph(
        'Click "Import OFL" to paste fixture JSON from open-fixture-library.org. '
        'The importer converts OFL\'s capability model to SlyLED format, handling '
        'multi-mode fixtures, 16-bit channels, and color detection automatically.'
    )
    add_screenshot('spa-settings-ofl-import.png', 'Open Fixture Library import modal')

    doc.add_heading(T('Import/Export Bundles'), level=2)
    doc.add_paragraph(
        'Export your custom profiles as a JSON bundle file for backup or sharing. '
        'Import bundles to add profiles from other installations.'
    )
    doc.add_page_break()

    # ── Section 9: Settings ───────────────────────────────────
    doc.add_heading(T('9. Settings & Configuration'), level=1)
    add_screenshot('spa-settings.png', 'Settings tab — app configuration and DMX output')
    doc.add_paragraph(
        'The Settings tab contains:\n'
        '- Orchestrator name and units\n'
        '- Stage dimensions\n'
        '- Dark mode toggle\n'
        '- Logging control\n'
        '- Configuration export/import\n'
        '- Show export/import\n'
        '- DMX output settings (Art-Net/sACN, frame rate, universe routing)\n'
        '- Fixture Profile Library\n'
        '- Factory reset'
    )

    doc.add_heading(T('Firmware'), level=2)
    add_screenshot('spa-firmware.png', 'Firmware tab — board detection and flash controls')
    doc.add_paragraph(
        'The Firmware tab provides:\n'
        '- COM port detection with board type identification\n'
        '- Serial version query\n'
        '- WiFi credential management\n'
        '- One-click firmware flashing for ESP32, D1 Mini, and Giga R1 boards'
    )
    doc.add_page_break()

    # ── Section 10: Camera Setup & Calibration ─────────────────
    doc.add_heading(T('10. Camera Setup & Calibration'), level=1)

    doc.add_heading(T('Adding Camera Nodes'), level=2)
    doc.add_paragraph(T(
        'Camera nodes are Orange Pi or Raspberry Pi SBCs running the SlyLED camera firmware. '
        'They capture snapshots for calibration, beam detection, depth estimation, and tracking.\n\n'
        'To add a camera:\n'
        '1. Connect the camera node to the same network as the orchestrator\n'
        '2. Go to Setup tab and click "Discover" — camera nodes respond to UDP broadcast\n'
        '3. Or manually add via Settings > Cameras with the camera\'s IP address\n'
        '4. Each USB camera sensor appears as a separate fixture that can be placed in the Layout'
    ))

    add_screenshot('spa-settings-cameras.png', 'Settings > Cameras — camera list with calibration status')

    doc.add_heading(T('ArUco Marker Calibration'), level=2)
    doc.add_paragraph(T(
        'ArUco markers are used for camera intrinsic calibration (lens parameters).\n\n'
        'Step-by-step:\n'
        '1. Print ArUco markers: Settings > Cameras > "Print ArUco Markers" (6 markers, 150mm each)\n'
        '2. Place printed markers on flat surfaces visible to the camera at various angles\n'
        '3. Click "Capture" 5+ times from different positions — each capture detects markers automatically\n'
        '4. Click "Compute" — the system runs cv2.calibrateCamera() to determine fx, fy, cx, cy\n'
        '5. RMS error < 2.0 indicates good calibration\n\n'
        'All ArUco detection runs on the orchestrator PC (not the camera hardware), so any camera '
        'that can serve JPEG snapshots works — even low-end Raspberry Pi boards.'
    ))

    doc.add_heading(T('Stage Map — Camera Pose'), level=2)
    doc.add_paragraph(T(
        'The Stage Map determines each camera\'s 3D position and orientation in stage coordinates '
        'using solvePnP with ArUco markers at known stage positions.\n\n'
        'Step-by-step:\n'
        '1. Place 3+ ArUco markers on the stage floor at known positions (measure with tape)\n'
        '2. Enter each marker\'s stage coordinates (X, Y, Z in mm) in the wizard\n'
        '3. Click "Compute Stage Map" — the system detects markers and solves for camera pose\n'
        '4. Result: camera position in stage mm, floor-plane homography, reprojection error\n\n'
        'The homography maps any pixel to stage floor coordinates (X, Y at Z=0), enabling '
        'beam detection in real-world millimeters rather than pixel space.'
    ))

    doc.add_heading(T('Mover Calibration Wizard'), level=2)
    doc.add_paragraph(T(
        'Moving head fixtures need calibration to map DMX pan/tilt values to real-world positions.\n\n'
        'The wizard runs automatically:\n'
        '1. Discovery: Spiral search to find the beam in the camera\'s field of view\n'
        '2. Adaptive settle: Waits 0.8-2.5s per move, verifies beam has stopped (double-capture)\n'
        '3. BFS mapping: Explores visible region, stops at boundaries when beam leaves camera view\n'
        '4. Grid build: Creates a bilinear interpolation grid for fast pan/tilt lookup\n'
        '5. Boundary verification: Flash on/off at edges to confirm true visibility limits\n\n'
        'After calibration, the system can aim any fixture at any stage coordinate using grid_inverse().'
    ))
    doc.add_page_break()

    # ── Section 11: 3D Environment Scanning ──────────────────
    doc.add_heading(T('11. 3D Environment Scanning'), level=1)

    doc.add_heading(T('Point Cloud Capture'), level=2)
    doc.add_paragraph(T(
        'The environment scan creates a 3D point cloud of your venue using depth estimation.\n\n'
        'Step-by-step:\n'
        '1. Position and calibrate 2+ cameras (Stage Map required)\n'
        '2. Go to Settings > Space > "Start Scan"\n'
        '3. Each camera captures a snapshot, the CV Engine runs Depth-Anything-V2 locally\n'
        '4. Point clouds from each camera are transformed to stage coordinates and merged\n'
        '5. Floor is auto-detected and normalized to Z=0\n\n'
        'Result: 5,000-20,000 colored 3D points representing your venue geometry.'
    ))

    doc.add_heading(T('Surface Detection'), level=2)
    doc.add_paragraph(T(
        'After scanning, click "Analyze Surfaces" to run RANSAC-based detection:\n\n'
        '- Floor: Horizontal plane detection (normal near vertical)\n'
        '- Walls: Vertical planes along stage boundaries\n'
        '- Obstacles: Clustered points that don\'t belong to floor or walls (pillars, furniture)\n\n'
        'Detected surfaces are used by the calibration system to predict where beams land '
        'and by the structured-light refinement to verify 3D accuracy.'
    ))

    doc.add_heading(T('CV Engine'), level=2)
    doc.add_paragraph(T(
        'All computer vision processing runs on the orchestrator PC, not the camera hardware. '
        'The CV Engine wraps three model pipelines:\n\n'
        '- Beam Detection: Fast color-filtered bright spot detection (<100ms)\n'
        '- Depth Estimation: Depth-Anything-V2 monocular depth (1-2s on x86)\n'
        '- Object Detection: YOLOv8n for person/object detection (100-200ms)\n\n'
        'Camera nodes only provide JPEG snapshots. This means any camera works — even a '
        'Raspberry Pi 3 that can\'t run ONNX models locally.'
    ))

    add_screenshot('spa-cv-status.png', 'CV Engine status — model loading state in Settings')
    doc.add_page_break()

    # ── Section 12: Stereo 3D & Light Mapping ─────────────────
    doc.add_heading(T('12. Stereo 3D & Light Mapping'), level=1)

    doc.add_heading(T('Stereo Triangulation'), level=2)
    doc.add_paragraph(T(
        'With 2+ calibrated cameras, the stereo engine triangulates 3D points from pixel '
        'observations. Each camera converts a pixel to a 3D ray; the intersection of rays '
        'from different cameras gives the 3D position in stage mm.\n\n'
        'Use cases:\n'
        '- Verify fixture positions by observing them from multiple cameras\n'
        '- Track moving objects (people) in 3D\n'
        '- Cross-validate depth estimation against triangulated ground truth\n\n'
        'Accuracy depends on camera baseline (distance between cameras) and calibration quality. '
        'Typical error: 20-100mm at 3-5m distance with 60\u00b0 FOV cameras.'
    ))

    doc.add_heading(T('Per-Fixture Light Maps'), level=2)
    doc.add_paragraph(T(
        'A light map records where each moving head\'s beam lands at every pan/tilt position '
        'in real stage coordinates.\n\n'
        'Step-by-step:\n'
        '1. Calibrate the fixture (mover calibration wizard)\n'
        '2. Click "Build Light Map" on the fixture\n'
        '3. The system sweeps the beam across the visible area in a grid pattern\n'
        '4. At each position: detect beam pixel, convert to stage mm via homography\n'
        '5. Result: (pan, tilt) \u2192 (stageX, stageY, stageZ) lookup table\n\n'
        'Inverse lookup: given a target position on stage, find the exact pan/tilt to aim there. '
        'Uses inverse-distance weighted interpolation of the 4 nearest samples.'
    ))

    doc.add_heading(T('Beam-as-Structured-Light'), level=2)
    doc.add_paragraph(T(
        'During calibration sweeps, each beam position that hits a surface provides a definite '
        '3D contact point. These points refine the environment model:\n\n'
        '- Floor height correction: Beam contacts adjust the detected floor Z value\n'
        '- Wall boundary extension: Contacts outside known wall boundaries expand the model\n'
        '- Surface verification: Predicted ray-surface intersections are compared with observed positions\n\n'
        'This is more accurate than depth estimation alone because the beam position is precise '
        '(bright centroid) and the ray direction is known (from pan/tilt geometry).'
    ))
    doc.add_page_break()

    # ── Section 13: Project Files ─────────────────────────────
    doc.add_heading(T('13. Project Files'), level=1)

    doc.add_heading(T('Export'), level=2)
    doc.add_paragraph(T(
        'File > Export saves your entire project as a .slyshow file containing:\n\n'
        '- All fixtures, children, layout positions, and stage dimensions\n'
        '- Actions, spatial effects, timelines, and show playlist\n'
        '- DMX settings, fixture profiles, and calibration data\n'
        '- Camera calibrations (intrinsics, stage maps, mover grids)\n'
        '- Point cloud (gzip-compressed for portability)\n'
        '- Light maps (per-fixture pan/tilt-to-stage lookup tables)\n\n'
        'The project file uses schema version 2. Files from older versions (v1) import '
        'successfully — new spatial fields default to empty.'
    ))

    doc.add_heading(T('Import'), level=2)
    doc.add_paragraph(T(
        'File > Import loads a .slyshow file, replacing all current state:\n\n'
        '1. All playback is stopped\n'
        '2. Fixtures, actions, timelines, and objects are restored\n'
        '3. Calibration data (homographies, grids, light maps) is restored\n'
        '4. Point cloud is decompressed and loaded into the 3D viewport\n'
        '5. Camera SSH passwords must be re-entered (not portable between machines)\n\n'
        'This enables a workflow where a designer builds the show on their laptop, exports the project, '
        'and a runtime operator loads it on the show PC — all spatial data transfers automatically.'
    ))

    doc.add_heading(T('Spatial Data in Projects'), level=2)
    doc.add_paragraph(T(
        'Point clouds can be 10,000-20,000 points (1-3 MB uncompressed). The export pipeline '
        'gzip-compresses the point array to ~200-500 KB. On import, the compressed data is '
        'automatically decompressed and loaded.\n\n'
        'Light maps (typically 200-300 samples per fixture) are stored alongside mover calibration data.'
    ))
    doc.add_page_break()

    # ── Section 14: System Limits ─────────────────────────────
    doc.add_heading(T('14. System Limits'), level=1)
    add_table(
        ['Resource', 'Limit', 'Notes'],
        [
            ['Children (performers)', '8 max', 'Protocol constant'],
            ['Strings per child', '8 max', 'ESP32 supports up to 8 GPIO pins'],
            ['LEDs per string', '65535 max', 'uint16_t addressing (protocol v4)'],
            ['Steps per runner', '16 max', 'LoadStepPayload array limit'],
            ['Bake segments per fixture', '64 max', 'Supports PT move time slices'],
            ['Bake frame rate', '40 Hz', 'Art-Net output rate'],
            ['DMX universes', 'Unlimited', 'One Art-Net packet per universe'],
            ['DMX channels per universe', '512', 'DMX-512 standard'],
            ['Pan/tilt resolution', '8 or 16 bit', 'Per profile channel definition'],
            ['Preview resolution', '1 fps', '1 color per string per second'],
            ['NTP sync offset', '5 seconds', 'GO command sent with future epoch'],
        ]
    )
    doc.add_page_break()

    # ── Section 15: Troubleshooting ──────────────────────────
    doc.add_heading(T('15. Troubleshooting'), level=1)

    problems = [
        ('3D Viewport Not Rendering', 'Browser doesn\'t support WebGL. Use Chrome, Firefox, or Edge.'),
        ('Performers Not Syncing', 'Children offline or on different network. Check Setup tab status.'),
        ('Bake Error: "No fixtures"', 'Add and position fixtures in Layout before baking.'),
        ('DMX Not Outputting', 'Check Art-Net engine is started (Settings → DMX Output → Start). '
         'Verify universe routing matches your bridge address.'),
        ('Moving Heads Not Tracking', 'Verify fixture has a profile with panRange/tiltRange > 0. '
         'Check rotation is set (Edit fixture → Pan/Tilt fields).'),
        ('Preview Shows No DMX Beams', 'Bake the timeline first. DMX preview requires baked data.'),
        ('Factory Reset', 'Settings → Factory Reset clears ALL data including fixtures, '
         'effects, timelines, and profiles.'),
    ]
    for title, desc in problems:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ── Section 16: API Quick Reference ──────────────────────
    doc.add_heading(T('16. API Quick Reference'), level=1)
    doc.add_paragraph('All endpoints are served from the Orchestrator at http://localhost:8080.')

    api_sections = [
        ('Fixtures', [
            ['GET/POST', '/api/fixtures', 'List / create'],
            ['GET/PUT/DELETE', '/api/fixtures/:id', 'CRUD'],
            ['PUT', '/api/fixtures/:id/aim', 'Set DMX aim point'],
            ['POST', '/api/fixtures/:id/resolve', 'Compute pixel positions'],
        ]),
        ('DMX Profiles', [
            ['GET/POST', '/api/dmx-profiles', 'List / create'],
            ['GET/PUT/DELETE', '/api/dmx-profiles/:id', 'CRUD'],
            ['GET', '/api/dmx-profiles/export', 'Export bundle'],
            ['POST', '/api/dmx-profiles/import', 'Import bundle'],
            ['POST', '/api/dmx-profiles/ofl/import-json', 'Import OFL fixture'],
        ]),
        ('Timelines', [
            ['GET/POST', '/api/timelines', 'List / create'],
            ['POST', '/api/timelines/:id/bake', 'Start baking'],
            ['POST', '/api/timelines/:id/start', 'Start playback'],
            ['POST', '/api/timelines/:id/stop', 'Stop playback'],
        ]),
        ('DMX Output', [
            ['GET', '/api/dmx/status', 'Engine status'],
            ['POST', '/api/dmx/start', 'Start engine'],
            ['POST', '/api/dmx/stop', 'Stop engine'],
            ['GET', '/api/dmx/fixture/:id/channels', 'Fixture channels with values'],
        ]),
        ('Calibration', [
            ['POST', '/api/cameras/:id/aruco/capture', 'Capture ArUco frame'],
            ['POST', '/api/cameras/:id/aruco/compute', 'Compute intrinsics'],
            ['POST', '/api/cameras/:id/stage-map', 'Compute camera pose'],
            ['POST', '/api/calibration/mover/:id/start', 'Start mover calibration'],
            ['GET', '/api/calibration/mover/:id/status', 'Poll calibration progress'],
            ['POST', '/api/calibration/stereo/calibrate', 'Build stereo engine'],
            ['POST', '/api/calibration/stereo/triangulate', 'Triangulate 3D point'],
        ]),
        ('CV Engine', [
            ['GET', '/api/cv/status', 'Model loading status'],
            ['POST', '/api/cameras/:id/detect', 'Object detection (local)'],
            ['POST', '/api/cameras/:id/depth', 'Depth estimation (local)'],
            ['POST', '/api/cameras/:id/beam-detect', 'Beam detection (local)'],
        ]),
        ('Space', [
            ['POST', '/api/space/scan', 'Start point cloud scan'],
            ['GET', '/api/space', 'Get point cloud'],
            ['POST', '/api/space/analyze', 'Detect surfaces'],
            ['GET', '/api/project/export', 'Export project (with spatial data)'],
            ['POST', '/api/project/import', 'Import project'],
        ]),
    ]
    for section_name, rows in api_sections:
        doc.add_heading(section_name, level=2)
        add_table(['Method', 'Endpoint', 'Description'], rows)
        doc.add_paragraph()

    # ── Save ───────────────────────────────────────────────────
    doc.save(OUTPUT)
    size = os.path.getsize(OUTPUT)
    print(f'Manual saved: {OUTPUT} ({size:,d} bytes)')
    return OUTPUT


def main():
    parser = argparse.ArgumentParser(description='Build SlyLED User Manual (Word doc)')
    parser.add_argument('--screenshots', action='store_true',
                        help='Regenerate screenshots before building')
    parser.add_argument('--lang', default='en', choices=['en', 'fr'],
                        help='Language: en (default) or fr')
    args = parser.parse_args()

    global _LANG, OUTPUT
    _LANG = args.lang
    if _LANG != 'en':
        OUTPUT = OUTPUT.replace('.docx', f'_{_LANG}.docx')

    if args.screenshots:
        print('Regenerating screenshots...')
        subprocess.run([sys.executable, os.path.join(os.path.dirname(__file__), 'screenshot_capture.py')],
                       check=True)
        print()

    print(f'Building Word document ({_LANG})...')
    path = build_manual()

    # Verify screenshots used
    used = 0
    missing = 0
    for f in os.listdir(SHOTS):
        if f.endswith('.png'):
            used += 1
    print(f'Screenshots used: {used} from {SHOTS}')
    print('Done.')


if __name__ == '__main__':
    main()
