#!/usr/bin/env python3
"""Convert docs/USER_MANUAL.md to .docx and .pdf.

Companion to tests/build_manual.py (which hand-crafts the styled Word manual
from scratch). This script renders the canonical markdown verbatim so newly
added sections — §20 Glossary (issue #663), Appendices A/B/C (#662) — show
up in the packaged outputs without waiting for build_manual.py to be ported.

Outputs land next to the markdown so they don't clobber
docs/SlyLED_User_Manual.docx / .pdf (the hand-crafted artefacts):

    docs/USER_MANUAL.docx
    docs/USER_MANUAL.pdf
    docs/USER_MANUAL_fr.docx     (with --fr)
    docs/USER_MANUAL_fr.pdf      (with --fr)

Usage:
    /usr/bin/python3 tests/build_manual_from_md.py
    /usr/bin/python3 tests/build_manual_from_md.py --fr
    /usr/bin/python3 tests/build_manual_from_md.py --no-pdf

Needs python-docx (for .docx) and Playwright Chromium (for .pdf). The
/usr/bin/python3 environment on this machine has both.
"""

from __future__ import annotations

import argparse
import base64
import html
import mimetypes
import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"


# ──────────────────────────────────────────────────────────────────────────
# Parsing
# ──────────────────────────────────────────────────────────────────────────

@dataclass
class Block:
    kind: str  # heading | para | code | table | list | rule | image | html | quote
    level: int = 0
    text: str = ""
    lang: str = ""
    rows: list[list[str]] = field(default_factory=list)
    items: list[tuple[int, bool, str]] = field(default_factory=list)  # (indent, ordered, text)
    caption: str = ""
    src: str = ""


def parse_markdown(md: str) -> list[Block]:
    lines = md.splitlines()
    blocks: list[Block] = []
    i = 0
    n = len(lines)

    def is_table_sep(s: str) -> bool:
        s = s.strip()
        if not s or "|" not in s:
            return False
        parts = [p.strip() for p in s.strip("|").split("|")]
        return bool(parts) and all(re.fullmatch(r":?-{3,}:?", p) for p in parts)

    def split_table_row(s: str) -> list[str]:
        # Token-aware split: protect pipes inside backtick code spans, link
        # text [...](...), and explicit \| escapes. Naive split() would shred
        # ``| `a \| b` | [x \| y](#z) |`` into 4 cells instead of 2.
        s = s.strip().strip("|")
        cells: list[str] = []
        buf: list[str] = []
        i = 0
        in_code = False
        in_link_text = False
        in_link_url = False
        while i < len(s):
            ch = s[i]
            if ch == "\\" and i + 1 < len(s) and s[i + 1] == "|":
                buf.append("|")
                i += 2
                continue
            if ch == "`" and not in_link_url:
                in_code = not in_code
                buf.append(ch)
                i += 1
                continue
            if not in_code:
                if ch == "[" and not in_link_text and not in_link_url:
                    in_link_text = True
                elif ch == "]" and in_link_text:
                    in_link_text = False
                    if i + 1 < len(s) and s[i + 1] == "(":
                        in_link_url = True
                        buf.append(ch)
                        buf.append("(")
                        i += 2
                        continue
                elif ch == ")" and in_link_url:
                    in_link_url = False
            if ch == "|" and not in_code and not in_link_text and not in_link_url:
                cells.append("".join(buf).strip())
                buf = []
                i += 1
                continue
            buf.append(ch)
            i += 1
        cells.append("".join(buf).strip())
        return cells

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Horizontal rule
        if re.fullmatch(r"-{3,}", stripped):
            blocks.append(Block("rule"))
            i += 1
            continue

        # Fenced code block
        m = re.match(r"^```(\w*)\s*$", stripped)
        if m:
            lang = m.group(1) or ""
            i += 1
            code_lines: list[str] = []
            while i < n and not re.match(r"^```\s*$", lines[i]):
                code_lines.append(lines[i])
                i += 1
            if i < n:
                i += 1  # closing fence
            blocks.append(Block("code", lang=lang, text="\n".join(code_lines)))
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.+?)\s*#*\s*$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            blocks.append(Block("heading", level=level, text=text))
            i += 1
            continue

        # Blockquote (possibly multi-line)
        if line.startswith(">"):
            quote_lines: list[str] = []
            while i < n and lines[i].startswith(">"):
                quote_lines.append(lines[i].lstrip(">").lstrip())
                i += 1
            blocks.append(Block("quote", text="\n".join(quote_lines)))
            continue

        # Table (header row + separator row + body)
        if "|" in stripped and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_table_row(stripped)
            i += 2  # skip header + separator
            body_rows: list[list[str]] = []
            while i < n and "|" in lines[i] and lines[i].strip():
                body_rows.append(split_table_row(lines[i]))
                i += 1
            blocks.append(Block("table", rows=[header] + body_rows))
            continue

        # List (ordered or unordered, possibly indented for nesting)
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)$", line)
        if m:
            items: list[tuple[int, bool, str]] = []
            while i < n:
                m2 = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)$", lines[i])
                if not m2:
                    # allow a blank line followed by more items
                    if lines[i].strip() == "" and i + 1 < n and re.match(
                        r"^(\s*)([-*]|\d+\.)\s+", lines[i + 1]
                    ):
                        i += 1
                        continue
                    break
                indent = len(m2.group(1)) // 2
                ordered = m2.group(2).endswith(".")
                text = m2.group(3)
                items.append((indent, ordered, text))
                i += 1
            blocks.append(Block("list", items=items))
            continue

        # Image (alone on a line)
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if m:
            alt, src = m.group(1), m.group(2)
            blocks.append(Block("image", caption=alt, src=src))
            i += 1
            continue

        # HTML block (anchor tags, self-contained)
        if stripped.startswith("<") and stripped.endswith(">") and "</" not in stripped:
            blocks.append(Block("html", text=stripped))
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Paragraph (accumulate until blank line or a block-starter)
        para_lines: list[str] = [line]
        i += 1
        while i < n:
            peek = lines[i]
            if not peek.strip():
                break
            if re.match(r"^#{1,6}\s", peek) or re.match(r"^```", peek) or peek.startswith(">"):
                break
            if re.match(r"^-{3,}\s*$", peek.strip()):
                break
            if re.match(r"^\s*([-*]|\d+\.)\s+", peek):
                break
            if "|" in peek and i + 1 < n and is_table_sep(lines[i + 1]):
                break
            if re.match(r"^!\[[^\]]*\]\([^)]+\)\s*$", peek.strip()):
                break
            para_lines.append(peek)
            i += 1
        blocks.append(Block("para", text=" ".join(ln.strip() for ln in para_lines)))

    return blocks


# ──────────────────────────────────────────────────────────────────────────
# Inline markdown → runs (for docx)
# ──────────────────────────────────────────────────────────────────────────

INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|"        # bold
    r"\*[^*]+\*|"              # italic
    r"__[^_]+__|"              # bold alt
    r"_[^_]+_|"                # italic alt
    r"`[^`]+`|"                # code
    r"\[[^\]]+\]\([^)]+\)|"    # link
    r"!\[[^\]]*\]\([^)]+\))"   # image-inline (rare)
)


def add_inline_runs(paragraph, text: str, base_bold: bool = False,
                    base_italic: bool = False, mono: bool = False) -> None:
    """Add runs to a docx paragraph with basic inline markdown handling."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif (part.startswith("_") and part.endswith("_")
                and not part.startswith("__") and len(part) > 2):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x1E, 0x4F, 0xC1)
                run.underline = True
        else:
            run = paragraph.add_run(part)
            if base_bold:
                run.bold = True
            if base_italic:
                run.italic = True
            if mono:
                run.font.name = "Consolas"


# ──────────────────────────────────────────────────────────────────────────
# docx renderer
# ──────────────────────────────────────────────────────────────────────────

def render_docx(blocks: list[Block], out_path: Path, title: str) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for h_idx, size in [(1, 22), (2, 16), (3, 13), (4, 12)]:
        hs = doc.styles[f"Heading {h_idx}"]
        hs.font.name = "Calibri"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1E, 0x4F, 0xC1) if h_idx == 1 else RGBColor(0x4B, 0x2E, 0xA8)

    for block in blocks:
        if block.kind == "heading":
            lvl = min(block.level, 4)
            p = doc.add_paragraph(style=f"Heading {lvl}")
            text = re.sub(r"<a id=\"[^\"]+\"></a>\s*", "", block.text)
            add_inline_runs(p, text)

        elif block.kind == "para":
            text = block.text.strip()
            if not text:
                continue
            # Skip raw anchor lines
            if re.fullmatch(r"<a id=\"[^\"]+\"></a>", text):
                continue
            p = doc.add_paragraph()
            add_inline_runs(p, text)

        elif block.kind == "code":
            lang = block.lang
            if lang == "mermaid":
                cap = doc.add_paragraph()
                cap_run = cap.add_run(f"[Mermaid diagram — see docs/diagrams/ for rendered source]")
                cap_run.italic = True
                cap_run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)

        elif block.kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline_runs(p, block.text.strip(), base_italic=True)

        elif block.kind == "rule":
            doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif block.kind == "list":
            for indent, ordered, text in block.items:
                style_name = "List Number" if ordered else "List Bullet"
                try:
                    p = doc.add_paragraph(style=style_name)
                except KeyError:
                    p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
                add_inline_runs(p, text)

        elif block.kind == "table":
            rows = block.rows
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncols)
            table.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell = table.rows[ri].cells[ci]
                    cell_text = row[ci] if ci < len(row) else ""
                    cell.paragraphs[0].text = ""
                    add_inline_runs(cell.paragraphs[0], cell_text, base_bold=(ri == 0))

        elif block.kind == "image":
            img_path = DOCS / block.src if not Path(block.src).is_absolute() else Path(block.src)
            # Markdown image paths are relative to the manual file
            candidates = [
                DOCS / block.src,
                ROOT / block.src,
            ]
            found = next((p for p in candidates if p.exists()), None)
            if found:
                try:
                    doc.add_picture(str(found), width=Inches(5.5))
                    if block.caption:
                        cap = doc.add_paragraph(block.caption)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cap.runs[0].italic = True
                except Exception as e:
                    doc.add_paragraph(f"[Image: {block.src} — failed to embed: {e}]")
            else:
                doc.add_paragraph(f"[Image: {block.src} — not found]")

        elif block.kind == "html":
            # Anchors are invisible in docx; skip silently
            if re.fullmatch(r'<a id="[^"]+"></a>', block.text):
                continue
            # Other HTML passes through as plain text
            p = doc.add_paragraph()
            run = p.add_run(block.text)
            run.font.color.rgb = RGBColor(0x6A, 0x6A, 0x6A)
            run.italic = True

    doc.save(str(out_path))


# ──────────────────────────────────────────────────────────────────────────
# HTML renderer (for PDF)
# ──────────────────────────────────────────────────────────────────────────

HTML_HEAD = """<!DOCTYPE html>
<html lang="{lang}">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
@page {{ size: Letter; margin: 0.75in; }}
body {{ font-family: Calibri, 'Segoe UI', Arial, sans-serif; font-size: 11pt; color: #222; line-height: 1.45; }}
h1 {{ color: #1E4FC1; font-size: 26pt; border-bottom: 2px solid #1E4FC1; padding-bottom: 4pt; margin-top: 18pt; }}
h2 {{ color: #4B2EA8; font-size: 18pt; margin-top: 14pt; border-bottom: 1px solid #CFCFCF; padding-bottom: 2pt; }}
h3 {{ color: #4B2EA8; font-size: 14pt; margin-top: 10pt; }}
h4 {{ color: #1E4FC1; font-size: 12pt; margin-top: 8pt; }}
code {{ font-family: Consolas, 'Courier New', monospace; background: #F3F3F3; padding: 1px 4px; border-radius: 2px; font-size: 10pt; }}
pre {{ font-family: Consolas, 'Courier New', monospace; background: #F3F3F3; padding: 10pt; border-left: 3px solid #1E4FC1; overflow-x: auto; font-size: 9pt; line-height: 1.3; page-break-inside: avoid; }}
pre.mermaid-src {{ border-left-color: #4B2EA8; }}
pre.mermaid-src::before {{ content: "[Mermaid diagram source — rendered version in docs/diagrams/]"; display: block; color: #6A6A6A; font-style: italic; font-size: 8.5pt; margin-bottom: 6pt; }}
table {{ border-collapse: collapse; margin: 8pt 0; width: 100%; font-size: 10pt; page-break-inside: avoid; }}
th {{ background: #EFEBFF; color: #2A1A6B; text-align: left; padding: 5pt 8pt; border: 1px solid #CFCFCF; }}
td {{ padding: 5pt 8pt; border: 1px solid #E4E4E4; vertical-align: top; }}
a {{ color: #1E4FC1; text-decoration: underline; }}
blockquote {{ border-left: 3px solid #FFB74D; background: #FFF6E3; margin: 8pt 0; padding: 8pt 12pt; font-style: italic; }}
blockquote strong {{ color: #8A4A00; font-style: normal; }}
hr {{ border: 0; border-top: 1px solid #CFCFCF; margin: 14pt 0; }}
ul, ol {{ margin: 6pt 0; padding-left: 20pt; }}
li {{ margin: 2pt 0; }}
img {{ max-width: 100%; }}
.emoji {{ font-family: "Segoe UI Emoji", "Apple Color Emoji", sans-serif; }}
</style>
</head>
<body>
"""

HTML_TAIL = "</body></html>"


def inline_to_html(text: str) -> str:
    # Escape first, then re-apply markup on the escaped version where safe.
    t = html.escape(text)
    # Link: [text](url) — apply before other transforms
    t = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', t)
    # Bold **x** and __x__
    t = re.sub(r"\*\*([^*]+)\*\*", r"<strong>\1</strong>", t)
    t = re.sub(r"__([^_]+)__", r"<strong>\1</strong>", t)
    # Italic *x* (avoid matching stand-alone asterisks in code)
    t = re.sub(r"(?<!\*)\*(?!\*)([^*\n]+)\*(?!\*)", r"<em>\1</em>", t)
    # Italic _x_ (avoid matching mid-word and __bold__ — require non-word
    # boundaries on both sides)
    t = re.sub(r"(?<![A-Za-z0-9_])_(?!_)([^_\n]+)_(?!_)(?![A-Za-z0-9_])",
                r"<em>\1</em>", t)
    # Inline code `x`
    t = re.sub(r"`([^`]+)`", r"<code>\1</code>", t)
    return t


def render_html(blocks: list[Block], title: str, lang: str = "en") -> str:
    parts: list[str] = [HTML_HEAD.format(title=html.escape(title), lang=lang)]

    for block in blocks:
        if block.kind == "heading":
            text = re.sub(r'<a id="([^"]+)"></a>\s*', "", block.text)
            slug = block.text.lower()
            slug = re.sub(r'<a id="([^"]+)"></a>', "", slug)
            slug = re.sub(r"[^\w\s-]", "", slug).strip().replace(" ", "-")
            lvl = min(block.level, 6)
            anchor_id = slug or f"h-{len(parts)}"
            parts.append(f'<h{lvl} id="{anchor_id}">{inline_to_html(text)}</h{lvl}>')

        elif block.kind == "para":
            text = block.text.strip()
            if not text or re.fullmatch(r'<a id="[^"]+"></a>', text):
                continue
            parts.append(f"<p>{inline_to_html(text)}</p>")

        elif block.kind == "code":
            cls = ' class="mermaid-src"' if block.lang == "mermaid" else ""
            parts.append(f"<pre{cls}><code>{html.escape(block.text)}</code></pre>")

        elif block.kind == "quote":
            parts.append(f"<blockquote>{inline_to_html(block.text.strip())}</blockquote>")

        elif block.kind == "rule":
            parts.append("<hr>")

        elif block.kind == "list":
            # Normalize into sequences of same-type siblings
            out: list[str] = []
            stack: list[tuple[int, bool]] = []  # (indent, ordered)
            for indent, ordered, text in block.items:
                while stack and stack[-1][0] > indent:
                    closing = "ol" if stack[-1][1] else "ul"
                    out.append(f"</{closing}>")
                    stack.pop()
                if not stack or stack[-1][0] < indent or stack[-1][1] != ordered:
                    opening = "ol" if ordered else "ul"
                    out.append(f"<{opening}>")
                    stack.append((indent, ordered))
                out.append(f"<li>{inline_to_html(text)}</li>")
            while stack:
                closing = "ol" if stack[-1][1] else "ul"
                out.append(f"</{closing}>")
                stack.pop()
            parts.append("".join(out))

        elif block.kind == "table":
            rows = block.rows
            if not rows:
                continue
            tparts = ["<table>"]
            header = rows[0]
            tparts.append("<thead><tr>")
            for h in header:
                tparts.append(f"<th>{inline_to_html(h)}</th>")
            tparts.append("</tr></thead>")
            tparts.append("<tbody>")
            for row in rows[1:]:
                tparts.append("<tr>")
                for cell in row:
                    tparts.append(f"<td>{inline_to_html(cell)}</td>")
                tparts.append("</tr>")
            tparts.append("</tbody></table>")
            parts.append("".join(tparts))

        elif block.kind == "image":
            candidates = [DOCS / block.src, ROOT / block.src]
            found = next((p for p in candidates if p.exists()), None)
            if found:
                # Inline as data: URI — Chromium loaded from a file:// page
                # sometimes blocks sibling file:// resources, and data URIs
                # also make the rendered HTML self-contained.
                mime = mimetypes.guess_type(str(found))[0] or "application/octet-stream"
                src_attr = f"data:{mime};base64,{base64.b64encode(found.read_bytes()).decode('ascii')}"
                parts.append(f'<img src="{src_attr}" alt="{html.escape(block.caption)}">')
                if block.caption:
                    parts.append(f"<p><em>{html.escape(block.caption)}</em></p>")
            else:
                parts.append(f"<p><em>[Image: {html.escape(block.src)} — not found]</em></p>")

        elif block.kind == "html":
            if re.fullmatch(r'<a id="[^"]+"></a>', block.text):
                m = re.match(r'<a id="([^"]+)"></a>', block.text)
                if m:
                    parts.append(f'<span id="{m.group(1)}"></span>')
            else:
                parts.append(block.text)

    parts.append(HTML_TAIL)
    return "\n".join(parts)


def render_pdf_via_playwright(html_path: Path, out_path: Path) -> None:
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        try:
            browser = p.chromium.launch()
        except Exception as exc:  # noqa: BLE001 — surface anything launch raises
            raise SystemExit(
                f"playwright chromium launch failed: {exc}\n"
                "Hint: run `playwright install chromium` once on this machine "
                "to install the browser binaries."
            ) from exc
        page = browser.new_page()
        page.goto(html_path.as_uri())
        page.wait_for_load_state("networkidle")
        # Margins are owned by the @page CSS rule; don't double-declare them
        # here (page.pdf would override CSS, which is confusing).
        page.pdf(
            path=str(out_path),
            format="Letter",
            print_background=True,
        )
        browser.close()


# ──────────────────────────────────────────────────────────────────────────
# Driver
# ──────────────────────────────────────────────────────────────────────────

def lint_known_limitations(md: str, md_path: Path) -> None:
    """Reject markdown constructs we know the parser cannot handle.

    Currently a single check (#677 finding #3): a blockquote (``> ...``)
    immediately followed by a fenced code block (`` ``` ``) is not recursed
    into. The fence is flattened into the quote text. The Pandoc-driven
    pipeline in ``tools/docs/build.py`` (#665) handles this correctly; this
    legacy renderer does not. Until the legacy fallback is retired, fail
    loudly if a writer adds the construct.
    """
    pattern = re.compile(r"(?m)^> .*\n(?:> .*\n)*> ```")
    m = pattern.search(md)
    if m:
        snippet = md[m.start(): m.start() + 120]
        raise SystemExit(
            f"lint: {md_path}: blockquote with nested code fence is not "
            f"supported by this legacy renderer (#677 limitation). "
            f"Move the fence outside the quote, or build via "
            f"tools/docs/build.py (pandoc).\n"
            f"  near: {snippet!r}"
        )


def build_one(md_path: Path, title: str, lang: str, make_pdf: bool) -> None:
    print(f"→ {md_path.name}")
    md = md_path.read_text(encoding="utf-8")
    lint_known_limitations(md, md_path)
    blocks = parse_markdown(md)
    print(f"  parsed {len(blocks)} blocks")

    docx_path = md_path.with_suffix(".docx")
    render_docx(blocks, docx_path, title)
    print(f"  wrote {docx_path.relative_to(ROOT)}")

    if make_pdf:
        html_str = render_html(blocks, title, lang=lang)
        html_path = md_path.with_suffix(".html")
        html_path.write_text(html_str, encoding="utf-8")
        pdf_path = md_path.with_suffix(".pdf")
        render_pdf_via_playwright(html_path, pdf_path)
        html_path.unlink()  # discard intermediate
        print(f"  wrote {pdf_path.relative_to(ROOT)}")


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0] if __doc__ else "")
    ap.add_argument("--fr", action="store_true", help="also build the French version")
    ap.add_argument("--no-pdf", action="store_true", help="skip the PDF step")
    args = ap.parse_args()

    make_pdf = not args.no_pdf

    build_one(DOCS / "USER_MANUAL.md",
              title="SlyLED User Manual",
              lang="en",
              make_pdf=make_pdf)

    if args.fr:
        build_one(DOCS / "USER_MANUAL_fr.md",
                  title="SlyLED — Manuel d'utilisation",
                  lang="fr",
                  make_pdf=make_pdf)

    return 0


if __name__ == "__main__":
    sys.exit(main())
